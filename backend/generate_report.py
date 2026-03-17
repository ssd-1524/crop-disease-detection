from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Maize Crop Disease Detection & Severity Analysis using Deep Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Project Report\n\n')

    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('Agriculture forms the backbone of the Indian economy, yet crop yield is highly susceptible to foliage diseases. This project presents a novel dual-stage deep learning pipeline for Maize (Corn) crop disease classification and precise severity segmentation. By fine-tuning a customized MobileNetV2 architecture for multi-class identification (Blight, Common Rust, Gray Leaf Spot, and Healthy) and utilizing Meta’s Segment Anything Model 2 (SAM2) for granular pixel-wise severity quantification, the proposed web-based tool bridges the gap between state-of-the-art vision models and accessible precision agriculture.')
    
    doc.add_page_break()

    # Chapter 1
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 The Importance of Maize as a Crop in India', level=2)
    p = doc.add_paragraph()
    p.add_run('Maize (Zea mays L.) is globally recognized as a fundamental agricultural commodity and stands as the third most important cereal crop in India, trailing only behind rice and wheat. Cultivated over approximately 9.5 million hectares in the country, maize contributes significantly to food security, acting both as human food and a critical component for animal feed and poultry industries. However, Indian maize farmers face severe challenges resulting from volatile environmental conditions which foster destructive foliar diseases like Northern Corn Leaf Blight, Common Rust, and Gray Leaf Spot. The reduction in yield caused by late detection of these diseases creates profound economic setbacks for subsistence and commercial farmers alike.\n')
    
    doc.add_heading('1.2 Objectives', level=2)
    doc.add_paragraph('- To develop an accurate image classification model (MobileNetV2) capable of distinguishing between major Maize diseases.\n- To employ zero-shot foundational segmentation (SAM2) to calculate the exact percentage of foliage affected by the disease.\n- To engineer a highly responsive, end-to-end full-stack web application integrating these AI capabilities using FastAPI, Next.js, and Supabase.')

    # Chapter 2
    doc.add_heading('2. System Architecture & Methodology', level=1)
    doc.add_paragraph('The application utilizes a decoupled, modern architecture to separate heavy AI inference from the user interface. \n[PLACE VISUAL ARCHITECTURE SCHEMATIC HERE]\n\n')
    doc.add_heading('2.1 Backend Strategy (FastAPI & PyTorch)', level=2)
    doc.add_paragraph('The backend thrives on an asynchronous Python runtime (FastAPI) which tightly interfaces with PyTorch. It executes a dual-model pipeline: first verifying if an image contains a disease, and if validated, chaining it into the SAM2 predictor for geographical isolation of the disease.')
    
    doc.add_page_break()

    # Chapter 3
    doc.add_heading('3. Custom Classification Model: MobileNetV2', level=1)
    doc.add_heading('3.1 Architecture Overview', level=2)
    doc.add_paragraph('MobileNetV2 is designed for mobile and resource-constrained environments, leveraging Depthwise Separable Convolutions and Inverted Residual Blocks with linear bottlenecks. Standard convolutions are factored into a spatial depthwise convolution and a 1x1 pointwise convolution, drastically reducing computational cost.')
    
    doc.add_heading('3.2 Customizations and Preprocessing', level=2)
    doc.add_paragraph('We modified the final dense classifier to output probabilities for exactly four classes: Blight, Common_Rust, Gray_Leaf_Spot, and Healthy. Feature-extraction layers were frozen (requires_grad = False), while the classification head was initialized using Xavier Uniform optimization to mitigate the vanishing gradient problem. \n\nImage Preprocessing logic implemented:\n1. Resize to 256x256\n2. Center Crop to 224x224\n3. Normalization (ImageNet means [0.485, 0.456, 0.406] and std [0.229, 0.224, 0.225]).\nTo maximize inference accuracy, Test Time Augmentation (TTA) aggregates predictions across multiple stochastic variations of the image at runtime.')

    # Code snippet
    code_para1 = doc.add_paragraph(style='Intense Quote')
    code_para1.add_run('''# MobileNetV2 Head Customization\nin_features = self.model.classifier[1].in_features\nself.model.classifier = nn.Sequential(\n    nn.Dropout(p=dropout_rate),\n    nn.Linear(in_features, num_classes),\n)\nnn.init.xavier_uniform_(self.model.classifier[1].weight)''')

    doc.add_page_break()

    # Chapter 4
    doc.add_heading('4. Disease Segmentation & Severity: SAM 2', level=1)
    doc.add_heading('4.1 Segment Anything Model 2 (Hiera Large)', level=2)
    doc.add_paragraph('While Classification informs "what" disease it is, Segmentation informs "how much". SAM2 was selected for its exceptional zero-shot performance and temporal/spatial mask generation robustness outperforming traditional U-Net.')

    doc.add_heading('4.2 The Color-Algorithmic Logic for Pathogen Isolation', level=2)
    doc.add_paragraph('SAM2 outputs dense region masks, but to distinguish diseased pixels from healthy leaf matter within the leaf boundary, a hybrid HSV and CIELAB color-profiling heuristic was engineered targeting specific disease morphology:\n\n- Common Rust: Features bright orange raised pustules. The logic focuses on high saturation and red/orange wrapping hues (HSV and LAB).\n- Blight: Identified through tan, cigar-shaped necrotic zones relying on a distinct brown spectrum exclusion.\n- Gray Leaf Spot: Rectangular, highly desaturated gray/straw lesions. Our implementation incorporates a dynamic "Greenness Departure" metric. Instead of looking for a static color, the algorithm profiles the specific "healthy green" of the current leaf and mathematically identifies pixels that are structurally devoid of green hue and drastically de-saturated.')

    code_para2 = doc.add_paragraph(style='Intense Quote')
    code_para2.add_run('''# Adaptive greenness-departure detector for GLS
# Evaluates what "healthy green" looks like on THIS leaf
green_zone = cv2.inRange(hsv, np.array([35, 45, 40]), np.array([85, 255, 255]))
healthy_s_median = float(np.median(green_saturations))
s_threshold = healthy_s_median * sensitivity
low_sat = (s_ch.astype(np.float32) < s_threshold).astype(np.uint8) * 255''')
    
    doc.add_page_break()

    # Summary
    doc.add_paragraph('[THIS IS A CONDENSED TEMPLATE GENERATED FOR YOU. PLEASE EXPAND THIS OUTLINE BY INSERTING THE REQUIRED DATASET STATISTICS, LOSS GRAPHS, AND APPLICATION SCREENSHOTS FOUND IN YOUR README. TO REACH 60 PAGES, INCLUDE DETAILED LITERATURE REVIEWS AND LARGE DATA TABLES AS PLACEHOLDERS REQUIRE.]')
    
    # Save
    doc.save('Crop_Disease_Detection_Detailed_Report.docx')
    print("Report generated successfully: Crop_Disease_Detection_Detailed_Report.docx")

if __name__ == "__main__":
    main()
